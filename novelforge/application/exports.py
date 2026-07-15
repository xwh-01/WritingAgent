"""User-facing document exports."""

from __future__ import annotations

import re
from pathlib import Path

from novelforge.domain import Story
from novelforge.storage.artifacts import ArtifactStore


class StoryExportService:
    def __init__(self, artifacts: ArtifactStore) -> None:
        self.artifacts = artifacts

    @staticmethod
    def safe_filename(title: str, fallback: str = "untitled") -> str:
        cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", title).strip(" ._")
        return cleaned[:80] or fallback

    def export_markdown(
        self,
        story: Story,
        output_path: str | Path | None = None,
    ) -> Path:
        output = (
            Path(output_path)
            if output_path
            else self.artifacts.path(
                story.id,
                "exports",
                f"{self.safe_filename(story.title)}.md",
            )
        )
        lines = [f"# {story.title}", "", f"> {story.premise}", ""]
        for index in sorted(story.manuscript.chapters):
            chapter = story.manuscript.chapters[index]
            if chapter.content.strip():
                lines.extend([f"## {chapter.title}", "", chapter.content, ""])
        output.write_text("\n".join(lines), encoding="utf-8")
        return output

    def export_docx(
        self,
        story: Story,
        output_path: str | Path | None = None,
    ) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        document = Document()
        document.styles["Normal"].font.name = "SimSun"
        document.styles["Normal"].font.size = Pt(12)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(story.title)
        title_run.bold = True
        title_run.font.size = Pt(22)

        if story.premise:
            premise = document.add_paragraph()
            premise.alignment = WD_ALIGN_PARAGRAPH.CENTER
            premise.add_run(story.premise).italic = True

        for index in sorted(story.manuscript.chapters):
            chapter = story.manuscript.chapters[index]
            if not chapter.content.strip():
                continue
            document.add_heading(chapter.title, level=1)
            for text in chapter.content.splitlines():
                paragraph = document.add_paragraph(text)
                paragraph.paragraph_format.first_line_indent = Cm(0.74)
                paragraph.paragraph_format.line_spacing = 1.5

        output = (
            Path(output_path)
            if output_path
            else self.artifacts.path(
                story.id,
                "exports",
                f"{self.safe_filename(story.title)}.docx",
            )
        )
        document.save(str(output))
        return output


__all__ = ["StoryExportService"]
