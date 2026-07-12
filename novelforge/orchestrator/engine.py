"""Finite-state workflow engine for NovelForge."""

from __future__ import annotations

import json
import re
from enum import StrEnum
from pathlib import Path
from typing import Callable
from uuid import UUID

from novelforge.agents import (
    CharacterArcAuditorAgent,
    ContinuityAuditorAgent,
    CriticAgent,
    EditorAgent,
    NovelDirectorAgent,
    PlannerAgent,
    SupervisorAgent,
    TaskEvaluatorAgent,
    WriterAgent,
)
from novelforge.application import AgentRunService, ContentService, MemoryService, QualityService
from novelforge.context.assembler import ContextAssembler
from novelforge.core.config import AppConfig, load_config
from novelforge.core.exceptions import PersistenceError, WorkflowError
from novelforge.core.models import (
    AgentTraceRun,
    AutoRevisionReport,
    AutonomousRunReport,
    BatchChapterResult,
    BatchWriteReport,
    Chapter,
    ChapterContract,
    CharacterFact,
    CharacterContinuityReport,
    ChapterOutline,
    ContinuityAuditReport,
    ReviewReport,
    RevisionProposal,
    Story,
    utc_now,
)
from novelforge.llm import build_llm_client
from novelforge.longform.manager import LongformManager
from novelforge.memory.graph_store import NetworkXGraphStore
from novelforge.memory.text_store import SQLiteFTSStore
from novelforge.memory.vector_store import ChromaVectorStore
from novelforge.orchestrator.auto_revisor import AutoRevisor, AutoRevisorConfig
from novelforge.orchestrator.trace_exporter import write_debug_report, write_trace_json
from novelforge.orchestrator.tool_registry import ToolRegistry
from novelforge.storage.repository import StoryRepository
from novelforge.validation import ChapterContractValidator


class WorkflowState(StrEnum):
    """小说写作工作流的有限状态枚举，覆盖从规划到完成的各阶段。"""

    PLANNING = "planning"
    OUTLINE_GENERATED = "outline_generated"
    CHAPTER_BEATS_READY = "chapter_beats_ready"
    CHAPTER_DRAFT = "chapter_draft"
    REVIEWING = "reviewing"
    REVISING = "revising"
    CHAPTER_FINALIZED = "chapter_finalized"
    COMPLETED = "completed"


class NovelForgeEngine:
    """小说写作流程的核心编排引擎，管理故事生命周期、智能体调用、记忆系统和持久化。"""

    def __init__(self, config: AppConfig | None = None):
        """初始化引擎，创建 LLM 客户端、向量/图/文本存储、各子智能体、上下文装配器等。"""
        self.config = config or load_config()
        self.llm = build_llm_client(self.config.llm)
        self.vector_store = ChromaVectorStore(self.config.memory.persist_directory)
        self.graph_store = NetworkXGraphStore(self.config.memory.graph_directory)
        self.text_store = SQLiteFTSStore(self.config.memory.sqlite_path)
        ranker_config = self.config.memory_ranker
        self.longform_manager = LongformManager(self.llm, memory_ranker_config=ranker_config)
        self.context_assembler = ContextAssembler(
            self.vector_store,
            self.graph_store,
            self.text_store,
            self.config.story.max_context_tokens,
            self.longform_manager,
            memory_ranker_config=ranker_config,
        )
        self.planner = PlannerAgent(self.llm)
        self.supervisor = SupervisorAgent(self.llm)
        self.director = NovelDirectorAgent(self.llm)
        self.task_evaluator = TaskEvaluatorAgent(self.llm)
        self.writer = WriterAgent(self.llm)
        self.critic = CriticAgent(self.llm)
        self.editor = EditorAgent(self.llm)
        self.continuity_auditor = ContinuityAuditorAgent(self.llm)
        self.character_arc_auditor = CharacterArcAuditorAgent(self.llm)
        self.content_service = ContentService()
        self.memory_service = MemoryService()
        self.quality_service = QualityService()
        self.agent_run_service = AgentRunService()
        self.story: Story | None = None
        self.last_review: dict[int, ReviewReport] = {}
        self.current_auto_revisor: AutoRevisor | None = None
        self.auto_status: str = "idle"
        self.repository = StoryRepository(
            database_path=self.config.storage.database_path,
            artifact_directory=self.config.storage.artifact_directory,
            legacy_state_directory=self.config.storage.legacy_state_directory,
        )

    @property
    def state_dir(self) -> Path:
        """Artifact directory for exports and debug files; canonical state resides in SQLite."""
        path = Path(self.config.storage.artifact_directory)
        path.mkdir(parents=True, exist_ok=True)
        return path

    def start_new_story(
        self,
        premise: str,
        title: str = "Untitled Novel",
        genre: str = "novel",
        style_guide: str = "",
    ) -> Story:
        """基于给定的前提、标题、体裁和文风指南创建并保存一个新故事。"""
        self.story = Story(title=title, premise=premise, genre=genre, style_guide=style_guide)
        self.save_state()
        return self.story

    def generate_outline(self, num_chapters: int | None = None, force: bool = False) -> list:
        """生成或补全故事大纲。若 force=True 则从零重新生成，否则只补足缺失的章节数。"""
        story = self._require_story()
        target_count = num_chapters or self.config.story.default_chapters
        if force:
            self.content_service.set_outlines(story, self.planner.generate_outline(story.premise, target_count))
        else:
            self._append_missing_outlines(story, target_count)
        story.status = WorkflowState.OUTLINE_GENERATED.value
        story.touch()
        self.save_state()
        return story.outlines

    def generate_beats(self, chapter_index: int) -> Chapter:
        """为指定章节生成场景节拍（scene beats），装配写作上下文后调用 Planner 代理。"""
        story = self._require_story()
        outline = story.get_outline(chapter_index)
        self.ensure_chapter_contract(chapter_index)
        context = self.context_assembler.assemble_writing_context(chapter_index, story)
        beats = self.planner.generate_beats(outline, context)
        chapter = story.chapters.get(chapter_index) or Chapter(index=chapter_index, title=outline.title)
        chapter.beats = beats
        self.content_service.save_chapter(story, chapter)
        story.status = WorkflowState.CHAPTER_BEATS_READY.value
        story.touch()
        self.save_state()
        return chapter

    def write_chapter(self, chapter_index: int) -> Chapter:
        """撰写指定章节的草稿。若尚未生成节拍则先调用 generate_beats，然后调用 Writer 代理写作，必要时润色。"""
        story = self._require_story()
        outline = story.get_outline(chapter_index)
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.beats:
            chapter = self.generate_beats(chapter_index)
        context = self.context_assembler.assemble_writing_context(chapter_index, story)
        contract = self.ensure_chapter_contract(chapter_index)
        content = self.writer.write_chapter(
            chapter_index, outline, chapter.beats, context, story.style_guide, contract=contract
        )
        content = self._polish_draft_if_enabled(story, chapter_index, content)
        chapter.update_content(content, status="draft", summary=outline.summary)
        self.content_service.save_chapter(story, chapter)
        story.status = WorkflowState.CHAPTER_DRAFT.value
        self._process_chapter_memory(story, chapter)
        story.touch()
        self.save_state()
        return chapter

    def ensure_chapter_contract(self, chapter_index: int, force: bool = False) -> ChapterContract:
        """返回章节合同；缺失时从大纲生成，force=True 时重新生成。"""
        story = self._require_story()
        existing = story.chapter_contracts.get(chapter_index)
        if existing is not None and not force:
            return existing
        outline = story.get_outline(chapter_index)
        contract = self.planner.generate_chapter_contract(story, outline)
        self.content_service.save_contract(story, contract)
        story.touch()
        self.save_state()
        return contract

    def update_chapter_contract(self, chapter_index: int, contract: ChapterContract) -> ChapterContract:
        """保存用户确认或编辑后的章节合同。"""
        story = self._require_story()
        story.get_outline(chapter_index)
        contract.chapter_index = chapter_index
        self.content_service.save_contract(story, contract)
        story.touch()
        self.save_state()
        return contract

    def list_character_facts(self, chapter_index: int | None = None) -> list[CharacterFact]:
        """列出全部事实，或列出指定章节生效的事实。"""
        story = self._require_story()
        if chapter_index is None:
            return list(story.character_facts)
        return self.longform_manager.fact_ledger.facts_at(story, chapter_index)

    def upsert_character_fact(self, fact: CharacterFact) -> CharacterFact:
        """保存用户确认的人物事实并持久化。"""
        story = self._require_story()
        saved = self.memory_service.confirm_fact(story, fact, self.longform_manager.fact_ledger)
        story.touch()
        self.save_state()
        return saved

    def delete_character_fact(self, fact_id: str) -> bool:
        """删除一条用户确认事实；系统提取事实不可直接删除。"""
        story = self._require_story()
        deleted = self.memory_service.remove_confirmed_fact(story, fact_id, self.longform_manager.fact_ledger)
        if deleted:
            story.touch()
            self.save_state()
        return deleted

    def _polish_draft_if_enabled(self, story: Story, chapter_index: int, content: str) -> str:
        """若配置启用了自动润色，用 Editor 代理对草稿进行小说质感提升后返回；否则直接返回原文。"""
        if not self.config.story.auto_polish_drafts:
            return content
        outline = story.get_outline(chapter_index)
        self.ensure_chapter_contract(chapter_index)
        instructions = (
            f"目标字数约 {self.config.story.prose_target_words} 字；"
            "把草稿改成更有小说质感的完整正文。"
            "保留章节大纲、节拍和长篇记忆中的事实，不改变结局走向。"
            "加强场景细节、人物动作、心理波动、对话潜台词和段落节奏。"
            f"本章标题: {outline.title}；核心冲突: {outline.conflict}；"
            f"文风: {story.style_guide or '清晰克制，有画面感，避免流水账。'}"
        )
        polished = self.editor.polish_prose(content, instructions)
        return polished or content

    def request_review(self, chapter_index: int) -> ReviewReport:
        """对指定章节执行评审：从向量存储召回记忆、获取长篇上下文，调用 Critic 代理生成评审报告。"""
        story = self._require_story()
        outline = story.get_outline(chapter_index)
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.content:
            raise WorkflowError(f"Chapter {chapter_index} has no draft to review.")
        memories = self.vector_store.query(
            "plot_summaries", outline.summary, k=5, story_id=str(story.id), max_chapter=chapter_index
        )
        memories.extend(self.vector_store.query(
            "memory_cards", outline.summary, k=5, story_id=str(story.id), max_chapter=chapter_index
        ))
        longform_context = self.longform_manager.get_enhanced_context(chapter_index, story, query=outline.summary)
        report = self.critic.review_chapter(
            chapter.content,
            outline,
            list(story.characters.values()),
            memories,
            longform_context,
        )
        checks = self.longform_manager.review_chapter_consistency(story, chapter_index, chapter.content)
        report.logic_issues.extend(checks["foreshadowing_issues"])
        report.pacing_issues.extend(checks["pacing_issues"])
        report.character_issues.extend(checks["character_state_issues"])
        self.last_review[chapter_index] = report
        chapter.status = "reviewed"
        story.status = WorkflowState.REVIEWING.value
        story.touch()
        self.save_state()
        return report

    def validate_chapter_contract(self, chapter_index: int) -> list:
        """对章节合同执行规则与语义联合验收，返回逐项证据。"""
        story = self._require_story()
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.content:
            raise WorkflowError(f"Chapter {chapter_index} has no content to validate.")
        contract = self.ensure_chapter_contract(chapter_index)
        validator = ChapterContractValidator(self.llm)
        return validator.validate(chapter.content, contract)

    def audit_chapter_continuity(self, chapter_index: int) -> ContinuityAuditReport:
        """对指定章节进行连续性审计，检查长篇小说的一致性，返回风险评分和审计报告。"""
        story = self._require_story()
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.content:
            raise WorkflowError(f"Chapter {chapter_index} has no content to audit.")
        outline = None
        try:
            outline = story.get_outline(chapter_index)
        except KeyError:
            outline = None
        query = outline.summary if outline else chapter.summary or chapter.title
        longform_context = self.longform_manager.get_enhanced_context(chapter_index, story, query=query)
        report = self.continuity_auditor.audit_chapter(story, chapter_index, chapter.content, longform_context)
        story.continuity_reports[chapter_index] = report
        story.touch()
        self.save_state()
        return report

    def audit_character_continuity(
        self,
        character_query: str,
        start_chapter: int,
        end_chapter: int,
    ) -> CharacterContinuityReport:
        """审计指定角色跨章节的人设与状态演变，并保存供 Director 后续修订使用的报告。"""
        story = self._require_story()
        if end_chapter < start_chapter:
            raise WorkflowError("end_chapter must be >= start_chapter")
        query = character_query.strip().lower()
        matches = [
            character for character in story.characters.values()
            if query in {character.id.lower(), character.name.lower()}
        ]
        if not matches:
            raise WorkflowError(f"Character '{character_query}' was not found in this story.")
        if len(matches) > 1:
            raise WorkflowError(f"Character '{character_query}' is ambiguous; use an exact character id.")
        report = self.character_arc_auditor.audit(
            story, matches[0], start_chapter, end_chapter
        )
        story.character_continuity_reports = [
            item for item in story.character_continuity_reports
            if not (
                item.character_id == report.character_id
                and item.start_chapter == start_chapter
                and item.end_chapter == end_chapter
            )
        ]
        story.character_continuity_reports.append(report)
        story.touch()
        self.save_state()
        return report

    def apply_revision(
        self,
        chapter_index: int,
        revised_content: str | None = None,
        revision_instruction: str = "",
    ) -> Chapter:
        """对指定章节应用修订：若未提供 revised_content 则用 Editor 代理基于评审报告自动修订。"""
        story = self._require_story()
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.content:
            raise WorkflowError(f"Chapter {chapter_index} has no content to revise.")
        report = self.last_review.get(chapter_index) or self.request_review(chapter_index)
        content = revised_content or self.editor.revise_chapter(
            chapter.content,
            report,
            story.style_guide,
            revision_instruction=revision_instruction,
        )
        chapter.update_content(content, status="revised", summary=chapter.summary)
        story.status = WorkflowState.REVISING.value
        story.touch()
        self._process_chapter_memory(story, chapter)
        self.save_state()
        return chapter

    def create_revision_proposal(self, chapter_index: int, instruction: str) -> RevisionProposal:
        """生成并验收修订候选，但不覆盖正式章节正文。"""
        story = self._require_story()
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.content:
            raise WorkflowError(f"Chapter {chapter_index} has no content to revise.")
        clean_instruction = instruction.strip()
        if not clean_instruction:
            raise WorkflowError("Revision instruction cannot be empty.")
        outline = story.get_outline(chapter_index)
        review = self.last_review.get(chapter_index) or self.critic.review_chapter(
            chapter.content,
            outline,
            list(story.characters.values()),
            [],
            "",
        )
        candidate = self.editor.revise_chapter(
            chapter.content,
            review,
            story.style_guide,
            revision_instruction=clean_instruction,
        )
        validation = self.critic.review_chapter(
            candidate,
            outline,
            list(story.characters.values()),
            [],
            "",
        )
        proposal = RevisionProposal(
            story_id=str(story.id),
            chapter_index=chapter_index,
            source_version=chapter.version,
            instruction=clean_instruction,
            original_content=chapter.content,
            proposed_content=candidate,
            review_report=review,
            validation_report=validation,
        )
        self.quality_service.add_proposal(story, proposal)
        story.touch()
        self.save_state()
        return proposal

    def get_revision_proposal(self, proposal_id: str) -> RevisionProposal | None:
        """查找当前故事中的修订候选。"""
        return self.quality_service.get_proposal(self._require_story(), proposal_id)

    def accept_revision_proposal(self, proposal_id: str) -> Chapter:
        """用户批准候选后，将其应用到仍处于源版本的正式章节。"""
        story = self._require_story()
        proposal = self.get_revision_proposal(proposal_id)
        if proposal is None:
            raise WorkflowError(f"Revision proposal not found: {proposal_id}")
        if proposal.status != "awaiting_approval":
            raise WorkflowError(f"Revision proposal is already {proposal.status}.")
        chapter = story.chapters.get(proposal.chapter_index)
        if chapter is None or chapter.version != proposal.source_version:
            raise WorkflowError("Chapter changed after this proposal was created; generate a new proposal.")
        chapter.update_content(proposal.proposed_content, status="revised", summary=chapter.summary)
        proposal.status = "accepted"
        proposal.updated_at = utc_now()
        story.status = WorkflowState.REVISING.value
        self._process_chapter_memory(story, chapter)
        story.touch()
        self.save_state()
        for run in story.agent_trace_runs:
            if proposal_id in run.proposal_ids and run.status == "awaiting_approval":
                run.status = "paused"
                if run.plan is not None:
                    run.plan.status = "paused"
                self.continue_director_agent(run.id, max_steps=6)
                break
        return chapter

    def reject_revision_proposal(self, proposal_id: str) -> RevisionProposal:
        """拒绝候选，不修改章节正文。"""
        proposal = self.get_revision_proposal(proposal_id)
        if proposal is None:
            raise WorkflowError(f"Revision proposal not found: {proposal_id}")
        if proposal.status != "awaiting_approval":
            raise WorkflowError(f"Revision proposal is already {proposal.status}.")
        proposal.status = "rejected"
        proposal.updated_at = utc_now()
        story = self._require_story()
        for run in story.agent_trace_runs:
            if proposal_id in run.proposal_ids and run.status == "awaiting_approval":
                run.status = "rejected"
                run.final_summary = "User rejected the revision proposal; the official chapter was unchanged."
                if run.plan is not None:
                    run.plan.status = "rejected"
        story.touch()
        self.save_state()
        return proposal

    def revise_revision_proposal(self, proposal_id: str, instruction: str) -> RevisionProposal:
        """拒绝旧候选，并结合用户反馈生成新的候选。"""
        old = self.get_revision_proposal(proposal_id)
        if old is None:
            raise WorkflowError(f"Revision proposal not found: {proposal_id}")
        feedback = instruction.strip()
        if not feedback:
            raise WorkflowError("Revision feedback cannot be empty.")
        story = self._require_story()
        linked_runs = [run for run in story.agent_trace_runs if proposal_id in run.proposal_ids]
        combined = f"{old.instruction}\n用户追加要求：{feedback}"
        self.reject_revision_proposal(proposal_id)
        new_proposal = self.create_revision_proposal(old.chapter_index, combined)
        for run in linked_runs:
            if new_proposal.id not in run.proposal_ids:
                run.proposal_ids.append(new_proposal.id)
            run.status = "awaiting_approval"
            run.final_summary = f"Created revised proposal {new_proposal.id}; waiting for approval."
            if run.plan is not None:
                run.plan.status = "awaiting_approval"
        story.touch()
        self.save_state()
        return new_proposal

    def update_chapter_content(
        self,
        chapter_index: int,
        content: str,
        title: str | None = None,
        status: str = "draft",
    ) -> Chapter:
        """手动更新章节内容、标题和状态，触发记忆处理后持久化。"""
        story = self._require_story()
        outline = None
        try:
            outline = story.get_outline(chapter_index)
        except KeyError:
            outline = None
        chapter = story.chapters.get(chapter_index) or Chapter(
            index=chapter_index,
            title=title or (outline.title if outline else f"Chapter {chapter_index}"),
        )
        if title:
            chapter.title = title
        chapter.update_content(content, status=status, summary=chapter.summary or (outline.summary if outline else ""))
        story.chapters[chapter_index] = chapter
        story.current_chapter = chapter_index
        story.status = WorkflowState.CHAPTER_DRAFT.value if status == "draft" else story.status
        self._process_chapter_memory(story, chapter)
        story.touch()
        self.save_state()
        return chapter

    def auto_write_chapter(self, chapter_index: int) -> AutoRevisionReport:
        """对指定章节启动自动写作→评审→修订循环（AutoRevisor），直至通过质量阈值或达到最大轮次。

        在启动 AutoRevisor 之前先运行连续性审计，将审计发现的问题注入修订循环，
        确保连续性问题和质量问题是同一条修复路径。
        """
        story = self._require_story()
        outline = story.get_outline(chapter_index)
        self.ensure_chapter_contract(chapter_index)
        chapter = story.chapters.get(chapter_index)
        if chapter is None or not chapter.beats:
            chapter = self.generate_beats(chapter_index)

        # Run continuity audit BEFORE the revision loop so findings can be fixed
        continuity_issues = []
        if chapter.content:
            try:
                audit_report = self.audit_chapter_continuity(chapter_index)
                continuity_issues = [
                    {"dimension": issue.dimension, "severity": issue.severity,
                     "description": issue.description, "evidence": issue.evidence,
                     "suggestion": issue.suggestion}
                    for issue in audit_report.issues
                ]
            except Exception:
                pass  # continuity audit is advisory; failure doesn't block writing

        config = AutoRevisorConfig(
            max_rounds=self.config.auto_revisor.max_rounds,
            pass_threshold=self.config.auto_revisor.pass_threshold,
            score_samples=self.config.auto_revisor.score_samples,
            quality_weights=self.config.auto_revisor.quality_weights,
        )
        self.current_auto_revisor = AutoRevisor(
            story=story,
            writer=self.writer,
            critic=self.critic,
            editor=self.editor,
            assembler=self.context_assembler,
            config=config,
        )
        self.auto_status = "running"
        result = self.current_auto_revisor.run(chapter_index, continuity_issues=continuity_issues or None)

        chapter = story.chapters.get(chapter_index) or Chapter(index=chapter_index, title=outline.title)
        chapter.update_content(
            result.final_content,
            status="revised" if result.passed else "reviewed",
            summary=chapter.summary or outline.summary,
        )
        story.chapters[chapter_index] = chapter
        story.current_chapter = chapter_index
        story.auto_revision_reports[chapter_index] = result
        story.status = WorkflowState.CHAPTER_FINALIZED.value if result.passed else WorkflowState.REVISING.value
        self._process_chapter_memory(story, chapter)
        story.touch()
        self.save_state()
        self.auto_status = self.current_auto_revisor.status
        return result

    def batch_write_chapters(
        self,
        start_chapter: int,
        end_chapter: int,
        use_auto_revision: bool = True,
        progress_callback: Callable[[dict[str, object]], None] | None = None,
    ) -> BatchWriteReport:
        """批量写作指定范围章节，支持自动修订模式，通过 progress_callback 返回实时进度。"""
        if start_chapter < 1 or end_chapter < start_chapter:
            raise WorkflowError("Invalid chapter range.")
        story = self._require_story()
        total = end_chapter - start_chapter + 1

        def emit_progress(
            message: str,
            chapter_index: int | None = None,
            stage: str = "working",
            progress_current: int | None = None,
        ) -> None:
            if progress_callback is None:
                return
            progress_callback(
                {
                    "message": message,
                    "chapter_index": chapter_index,
                    "stage": stage,
                    "progress_current": progress_current if progress_current is not None else 0,
                    "progress_total": total,
                    "status": "running_batch",
                }
            )

        if len(story.outlines) < end_chapter:
            emit_progress(
                f"Generating outline up to chapter {end_chapter}",
                stage="outline",
                progress_current=0,
            )
        self._ensure_outlines(end_chapter)
        report = BatchWriteReport(
            start_chapter=start_chapter,
            end_chapter=end_chapter,
            use_auto_revision=use_auto_revision,
        )
        for chapter_index in range(start_chapter, end_chapter + 1):
            try:
                chapter = story.chapters.get(chapter_index)
                if chapter is None or not chapter.beats:
                    emit_progress(
                        f"Chapter {chapter_index}: generating beats",
                        chapter_index=chapter_index,
                        stage="beats",
                        progress_current=report.completed,
                    )
                    self.generate_beats(chapter_index)
                if use_auto_revision:
                    emit_progress(
                        f"Chapter {chapter_index}: writing, reviewing, and revising",
                        chapter_index=chapter_index,
                        stage="auto_revision",
                        progress_current=report.completed,
                    )
                    auto_report = self.auto_write_chapter(chapter_index)
                    chapter = story.chapters[chapter_index]
                    report.results.append(
                        BatchChapterResult(
                            chapter_index=chapter_index,
                            status="passed" if auto_report.passed else "reviewed",
                            title=chapter.title,
                            word_count=len(chapter.content),
                            auto_revision_score=auto_report.final_score,
                            message=f"{len(auto_report.rounds)} auto-revision rounds",
                        )
                    )
                else:
                    emit_progress(
                        f"Chapter {chapter_index}: writing draft",
                        chapter_index=chapter_index,
                        stage="draft",
                        progress_current=report.completed,
                    )
                    chapter = self.write_chapter(chapter_index)
                    report.results.append(
                        BatchChapterResult(
                            chapter_index=chapter_index,
                            status=chapter.status,
                            title=chapter.title,
                            word_count=len(chapter.content),
                            message="draft generated",
                        )
                    )
                report.completed += 1
                emit_progress(
                    f"Chapter {chapter_index}: completed",
                    chapter_index=chapter_index,
                    stage="completed",
                    progress_current=report.completed,
                )
            except Exception as exc:
                report.failed += 1
                emit_progress(
                    f"Chapter {chapter_index}: failed - {exc}",
                    chapter_index=chapter_index,
                    stage="failed",
                    progress_current=report.completed,
                )
                report.results.append(
                    BatchChapterResult(
                        chapter_index=chapter_index,
                        status="failed",
                        message=str(exc),
                    )
                )
        self.agent_run_service.add_batch_report(story, report)
        story.touch()
        self.save_state()
        return report

    def agentic_writing_run(
        self,
        objective: str,
        start_chapter: int,
        end_chapter: int,
        use_auto_revision: bool = True,
        progress_callback: Callable[[dict[str, object]], None] | None = None,
    ) -> AutonomousRunReport:
        """已弃用：保留兼容旧调用；新产品流程使用章节合同和批量写作。"""
        import warnings
        warnings.warn(
            "agentic_writing_run is deprecated; use chapter contracts with batch_write_chapters",
            DeprecationWarning,
            stacklevel=2,
        )
        if start_chapter < 1 or end_chapter < start_chapter:
            raise WorkflowError("Invalid chapter range.")
        story = self._require_story()
        run = self.supervisor.plan_writing_run(
            story=story,
            objective=objective,
            start_chapter=start_chapter,
            end_chapter=end_chapter,
            use_auto_revision=use_auto_revision,
        )
        self.agent_run_service.add_autonomous_run(story, run)
        run.status = "running"
        run.updated_at = utc_now()
        story.touch()
        self.save_state()
        self._emit_agent_progress(
            run,
            "SupervisorAgent",
            "plan_run",
            f"Planned with {run.planning_strategy}: {run.planning_notes or run.summary}",
            progress_callback,
        )

        for task in run.tasks:
            task.status = "running"
            task.started_at = utc_now()
            run.updated_at = utc_now()
            story.touch()
            self.save_state()
            self._emit_agent_progress(
                run,
                task.agent,
                task.action,
                task.reason,
                progress_callback,
                chapter_index=task.chapter_index,
                task_id=task.id,
            )
            try:
                task.output_summary = self._execute_agent_task(task, end_chapter)
                task.status = "completed"
                task.completed_at = utc_now()
                run.completed_tasks += 1
                self._emit_agent_progress(
                    run,
                    task.agent,
                    task.action,
                    task.output_summary,
                    progress_callback,
                    chapter_index=task.chapter_index,
                    task_id=task.id,
                )
            except Exception as exc:
                task.status = "failed"
                task.error = str(exc)
                task.completed_at = utc_now()
                run.failed_tasks += 1
                self._emit_agent_progress(
                    run,
                    task.agent,
                    task.action,
                    f"Failed: {exc}",
                    progress_callback,
                    chapter_index=task.chapter_index,
                    task_id=task.id,
                )
                break
            finally:
                run.updated_at = utc_now()
                story.touch()
                self.save_state()

        run.status = "failed" if run.failed_tasks else "completed"
        run.summary = (
            f"Agentic run {run.status}: {run.completed_tasks}/{len(run.tasks)} tasks completed, "
            f"{run.failed_tasks} failed."
        )
        run.updated_at = utc_now()
        story.touch()
        self.save_state()
        return run

    def run_director_agent(self, user_message: str, max_steps: int = 6) -> AgentTraceRun:
        """运行导演智能体（NovelDirector），根据用户指令调用已注册的工具完成多步操作。"""
        story = self._require_story()
        registry = ToolRegistry(self)
        run = self.director.run(
            story_id=str(story.id),
            user_message=user_message,
            max_steps=max_steps,
            story=story,
            tool_registry=registry,
            task_evaluator=self.task_evaluator,
        )
        self.agent_run_service.add_director_run(story, run)
        story.touch()
        self.save_state()
        return run

    def resume_director_agent(self, run_id: str, user_response: str, max_steps: int = 6) -> AgentTraceRun:
        """用用户对追问的回答恢复同一次 Director 运行。"""
        story = self._require_story()
        run = self.get_director_run(run_id)
        if run is None:
            raise WorkflowError(f"Director trace run not found: {run_id}")
        if run.status != "needs_user_input" or not run.pending_question:
            raise WorkflowError("Director run is not waiting for user input.")
        answer = user_response.strip()
        if not answer:
            raise WorkflowError("User response cannot be empty.")
        question = run.pending_question
        if run.pending_user_question is not None:
            run.pending_user_question.answer = answer
        run.user_responses.append(answer)
        run.pending_question = ""
        run.status = "running"
        run.final_summary = ""
        resumed_message = (
            f"原始任务：{run.user_message}\n"
            f"Director 追问：{question}\n"
            f"用户回答：{answer}\n"
            "请根据回答继续完成原始任务。"
        )
        registry = ToolRegistry(self)
        run.plan = self.director.create_plan(story, resumed_message, registry.list_specs())
        self.director.run(
            story_id=str(story.id),
            user_message=resumed_message,
            max_steps=max_steps,
            story=story,
            tool_registry=registry,
            task_evaluator=self.task_evaluator,
            existing_run=run,
        )
        story.touch()
        self.save_state()
        return run

    def continue_director_agent(self, run_id: str, max_steps: int = 6) -> AgentTraceRun:
        """从持久化检查点继续一个 paused Director 运行。"""
        story = self._require_story()
        run = self.get_director_run(run_id)
        if run is None:
            raise WorkflowError(f"Director trace run not found: {run_id}")
        if run.status not in {"paused", "running"}:
            raise WorkflowError(f"Director run cannot continue from status {run.status}.")
        registry = ToolRegistry(self)
        self.director.run(
            story_id=str(story.id),
            user_message=run.user_message,
            max_steps=max_steps,
            story=story,
            tool_registry=registry,
            task_evaluator=self.task_evaluator,
            existing_run=run,
        )
        story.touch()
        self.save_state()
        return run

    def list_director_runs(self) -> list[AgentTraceRun]:
        """列出当前故事的所有导演智能体运行记录。"""
        return list(self._require_story().agent_runs.director)

    def get_director_run(self, run_id: str) -> AgentTraceRun | None:
        """根据 run_id 获取单次导演智能体运行记录，不存在时返回 None。"""
        return self.agent_run_service.get_director_run(self._require_story(), run_id)

    def export_director_trace_json(self, run_id: str, output_path: str | Path | None = None) -> Path:
        """导出指定导演运行记录的轨迹 JSON 文件。"""
        run = self.get_director_run(run_id)
        if run is None:
            raise WorkflowError(f"Director trace run not found: {run_id}")
        story = self._require_story()
        output = Path(output_path) if output_path else self.repository.artifact_path(
            story.id, "traces", f"{run.id}.json"
        )
        return write_trace_json(run, output)

    def export_director_debug_report(self, run_id: str, output_path: str | Path | None = None) -> Path:
        """导出指定导演运行记录的调试 Markdown 报告。"""
        run = self.get_director_run(run_id)
        if run is None:
            raise WorkflowError(f"Director trace run not found: {run_id}")
        story = self._require_story()
        output = Path(output_path) if output_path else self.repository.artifact_path(
            story.id, "traces", f"{run.id}.debug.md"
        )
        return write_debug_report(run, output)

    def _execute_agent_task(self, task, end_chapter: int) -> str:
        """根据任务对象中的 action 字段执行对应的引擎操作（大纲、节拍、写作、审计、记忆等），返回结果摘要。"""
        story = self._require_story()
        if task.action == "ensure_outline":
            before = len(story.outlines)
            if before < end_chapter:
                self._ensure_outlines(end_chapter)
                return f"Generated outlines through chapter {end_chapter}."
            return f"Outline already covers {before} chapters."
        if task.chapter_index is None:
            raise WorkflowError(f"Task {task.id} requires a chapter index.")
        if task.action == "generate_beats":
            chapter = self.generate_beats(task.chapter_index)
            return f"Generated {len(chapter.beats)} scene beats."
        if task.action == "write_chapter":
            chapter = self.write_chapter(task.chapter_index)
            return f"Wrote draft with {len(chapter.content)} characters."
        if task.action == "auto_write_chapter":
            report = self.auto_write_chapter(task.chapter_index)
            return f"Auto-revision finished with score {report.final_score:.2f}; passed={report.passed}."
        if task.action == "audit_chapter_continuity":
            report = self.audit_chapter_continuity(task.chapter_index)
            return f"Continuity risk {report.risk_score:.1f}; passed={report.passed}."
        if task.action == "memory_checkpoint":
            chapter = story.chapters.get(task.chapter_index)
            if chapter is None or not chapter.content:
                raise WorkflowError(f"Chapter {task.chapter_index} has no content to index.")
            self._process_chapter_memory(story, chapter)
            return (
                f"Memory updated: {len(story.memory_cards)} cards, "
                f"{len(story.chapter_summaries)} chapter summaries."
            )
        raise WorkflowError(f"Unknown agent task action: {task.action}")

    def _emit_agent_progress(
        self,
        run: AutonomousRunReport,
        agent: str,
        action: str,
        message: str,
        progress_callback: Callable[[dict[str, object]], None] | None,
        chapter_index: int | None = None,
        task_id: str | None = None,
    ) -> None:
        """向进度回调发送智能体任务的当前状态和进度信息。"""
        if progress_callback is None:
            return
        progress_callback(
            {
                "status": "running_agentic" if run.status == "running" else run.status,
                "run_id": run.id,
                "task_id": task_id,
                "agent": agent,
                "action": action,
                "message": message,
                "chapter_index": chapter_index,
                "stage": action,
                "progress_current": run.completed_tasks,
                "progress_total": len(run.tasks),
                "planning_strategy": run.planning_strategy,
            }
        )

    def get_auto_status(self) -> dict[str, object]:
        """返回当前自动修订器的运行状态、轮次和停止请求标志。"""
        if self.current_auto_revisor is None:
            return {"status": self.auto_status, "round": 0, "stop_requested": False}
        return {
            "status": self.current_auto_revisor.status,
            "round": self.current_auto_revisor.current_round,
            "stop_requested": self.current_auto_revisor.stop_requested,
        }

    def stop_auto_revision(self) -> bool:
        """请求停止当前正在运行的自动修订循环，返回是否成功发送停止信号。"""
        if self.current_auto_revisor is None:
            return False
        self.current_auto_revisor.request_stop()
        self.auto_status = "stop_requested"
        return True

    def finalize_chapter(self, chapter_index: int) -> Chapter:
        """将章节标记为终稿状态，触发记忆处理，若为最后一章则标记整个故事已完成。"""
        story = self._require_story()
        chapter = story.chapters[chapter_index]
        chapter.status = "finalized"
        story.status = WorkflowState.CHAPTER_FINALIZED.value
        if chapter_index >= len(story.outlines):
            story.status = WorkflowState.COMPLETED.value
        if chapter.content:
            self._process_chapter_memory(story, chapter)
        story.touch()
        self.save_state()
        return chapter

    def advance_to_next_chapter(self) -> Chapter:
        """推进到下一章并生成节拍；若已是最后一章则标记故事完成并抛出异常。"""
        story = self._require_story()
        next_index = max(story.current_chapter + 1, 1)
        if next_index > len(story.outlines):
            story.status = WorkflowState.COMPLETED.value
            self.save_state()
            raise WorkflowError("Story is completed; no next chapter exists.")
        return self.generate_beats(next_index)

    def _ensure_outlines(self, end_chapter: int) -> None:
        """确保目纲至少覆盖到 end_chapter，不足则补生成并更新状态。"""
        story = self._require_story()
        self._append_missing_outlines(story, end_chapter)
        story.status = WorkflowState.OUTLINE_GENERATED.value
        story.touch()
        self.save_state()

    def _append_missing_outlines(self, story: Story, target_count: int) -> None:
        """若现有大纲数量不足 target_count，补生成缺失的章节并追加到故事大纲列表中。"""
        existing_count = len(story.outlines)
        if existing_count >= target_count:
            return
        missing_count = target_count - existing_count
        generated = self.planner.generate_outline(story.premise, missing_count)
        for offset, outline in enumerate(generated[:missing_count], start=1):
            real_index = existing_count + offset
            story.outlines.append(self._renumber_outline(outline, real_index))

    def _renumber_outline(self, outline: ChapterOutline, chapter_index: int) -> ChapterOutline:
        """重新编号大纲项，更改 chapter_index 为目标值后返回副本。"""
        return outline.model_copy(update={"chapter_index": chapter_index})

    def save_state(self) -> Path:
        """将当前故事状态持久化到仓库（repository），返回保存路径。"""
        story = self._require_story()
        try:
            return self.repository.save(story)
        except Exception as exc:
            raise PersistenceError(f"Could not save story state: {exc}") from exc

    def load_state(self, story_id: str | UUID) -> Story:
        """从仓库加载指定 story_id 的故事状态，若不存在或读取失败则抛出 PersistenceError。"""
        if not self.repository.exists(story_id):
            raise PersistenceError(f"Story state not found: {self.repository.story_path(story_id)}")
        try:
            self.story = self.repository.load(story_id)
            return self.story
        except Exception as exc:
            raise PersistenceError(f"Could not load story state: {exc}") from exc

    def export_markdown(self, output_path: str | Path | None = None) -> Path:
        """将当前故事的所有章节导出为单个 Markdown 文件。"""
        story = self._require_story()
        output = Path(output_path) if output_path else self.repository.artifact_path(
            story.id, "exports", f"{self._safe_export_filename(story.title)}.md"
        )
        lines = [f"# {story.title}", "", f"> {story.premise}", ""]
        for index in sorted(story.chapters):
            chapter = story.chapters[index]
            lines.extend([f"## {chapter.title}", "", chapter.content, ""])
        output.write_text("\n".join(lines), encoding="utf-8")
        return output

    def export_docx(self, output_path: str | Path | None = None) -> Path:
        """将当前故事的所有章节导出为 .docx 格式文档，含标题页和章节正文。"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        story = self._require_story()
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(12)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(story.title)
        run.bold = True
        run.font.size = Pt(22)

        if story.premise:
            premise_para = doc.add_paragraph()
            premise_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = premise_para.add_run(story.premise)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = None

        doc.add_paragraph()

        for idx in sorted(story.chapters):
            chapter = story.chapters[idx]
            if chapter.content:
                heading = doc.add_heading(chapter.title, level=1)
                for run in heading.runs:
                    run.font.size = Pt(16)
                for paragraph_text in chapter.content.split("\n"):
                    para = doc.add_paragraph(paragraph_text)
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    para.paragraph_format.line_spacing = 1.5

        output = Path(output_path) if output_path else self.repository.artifact_path(
            story.id, "exports", f"{self._safe_export_filename(story.title)}.docx"
        )
        doc.save(str(output))
        return output

    def _safe_export_filename(self, title: str, fallback: str = "untitled") -> str:
        """清理标题中的非法文件名字符，截断至 80 字符，返回安全的导出文件名。"""
        cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", title).strip(" ._")
        return cleaned[:80] or fallback

    def export_auto_revision_report(self, chapter_index: int, output_path: str | Path | None = None) -> Path:
        """导出指定章节的自动修订报告到文件，若不存在则抛出 WorkflowError。"""
        story = self._require_story()
        report = story.auto_revision_reports.get(chapter_index)
        if report is None:
            raise WorkflowError(f"No auto-revision report for chapter {chapter_index}.")
        return self.repository.export_auto_revision_report(story, report, output_path)

    def delete_story_data(self, story_id: str | UUID) -> dict[str, object]:
        """删除指定故事的所有数据：仓库文件、向量索引、全文索引和图节点。"""
        story_id_str = str(story_id)
        result = {
            "story_id": story_id_str,
            "story_file": self.repository.delete(story_id_str),
            "vector_items": self.vector_store.delete_story(story_id_str),
            "fts_items": self.text_store.delete_story(story_id_str),
            "graph_nodes": self.graph_store.delete_story(story_id_str),
        }
        if self.story and str(self.story.id) == story_id_str:
            self.story = None
        return result

    def rebuild_derived_indexes(self, story_id: str | UUID | None = None) -> dict[str, int | str]:
        """Rebuild all disposable indexes from canonical SQLite-backed story state.

        This operation deliberately reads only Story data and never treats an existing vector, FTS,
        or graph entry as authoritative.
        """
        if story_id is not None and (self.story is None or str(self.story.id) != str(story_id)):
            self.load_state(story_id)
        story = self._require_story()
        story_id_str = str(story.id)
        self.vector_store.delete_story(story_id_str)
        self.text_store.delete_story(story_id_str)
        self.graph_store.delete_story(story_id_str)
        indexed_chapters = 0
        for chapter in story.chapters.values():
            if chapter.content:
                self._index_chapter(story, chapter)
                indexed_chapters += 1
        cards = story.memory_cards
        if cards:
            self.vector_store.add(
                "memory_cards",
                [card.content for card in cards],
                [
                    {
                        "story_id": story_id_str,
                        "type": card.type,
                        "chapter": card.chapter,
                        "importance": card.importance,
                        "entities": ",".join(card.entities),
                        "tags": ",".join(card.tags),
                    }
                    for card in cards
                ],
                [card.id if card.id.startswith(f"{story_id_str}:") else f"{story_id_str}:memory_card:{card.id}" for card in cards],
            )
        characters = list(story.characters.values())
        if characters:
            self.vector_store.add(
                "characters",
                [
                    " ".join(filter(None, [item.name, item.appearance, item.personality, item.motivation, item.weakness, item.arc]))
                    for item in characters
                ],
                [{"story_id": story_id_str, "type": "character", "character_id": item.id} for item in characters],
                [f"{story_id_str}:character:{item.id}" for item in characters],
            )
            for character in characters:
                attrs = character.model_dump()
                attrs["story_id"] = story_id_str
                self.graph_store.add_node(f"{story_id_str}:character:{character.id}", attrs)
        world_settings = story.world_settings
        if world_settings:
            self.vector_store.add(
                "world",
                [item.content for item in world_settings],
                [{"story_id": story_id_str, "type": "world", "category": item.category, **item.metadata} for item in world_settings],
                [f"{story_id_str}:world:{item.id}" for item in world_settings],
            )
        pending = [
            int(event["id"]) for event in self.repository.pending_index_events()
            if str(event["story_id"]) == story_id_str
        ]
        self.repository.mark_index_events_processed(pending)
        return {
            "story_id": story_id_str,
            "chapters": indexed_chapters,
            "memory_cards": len(cards),
            "characters": len(characters),
            "world_settings": len(world_settings),
            "events_processed": len(pending),
        }

    def _index_chapter(self, story: Story, chapter: Chapter) -> None:
        """将章节内容写入全文索引和向量存储用于后续语义检索。"""
        prefix = f"{story.id}:chapter:{chapter.index}:"
        self.text_store.delete_prefix(prefix)
        self.vector_store.delete_prefix("plot_summaries", prefix)
        doc_id = f"{prefix}current"
        self.text_store.index_document(doc_id, chapter.content)
        self.vector_store.add(
            "plot_summaries",
            [chapter.summary or chapter.content[:500]],
            [{"story_id": str(story.id), "type": "chapter_summary", "chapter": chapter.index, "version": chapter.version}],
            [doc_id],
        )

    def _process_chapter_memory(self, story: Story, chapter: Chapter) -> None:
        """处理章节记忆：索引内容、调用长篇管理器提取角色/世界观，写入向量/图存储并审计连续性。"""
        self._index_chapter(story, chapter)
        result = self.longform_manager.process_new_chapter(story, chapter.index, chapter.content)
        extraction = result.get("extraction") if isinstance(result, dict) else None
        if extraction is not None:
            self._index_extracted_memory(story, extraction)
        memory = result.get("memory", {}) if isinstance(result, dict) else {}
        cards = memory.get("memory_cards", []) if isinstance(memory, dict) else []
        if not cards:
            self._audit_processed_chapter(story, chapter)
            return
        self.vector_store.add(
            "memory_cards",
            [card.content for card in cards],
            [
                {
                    "story_id": str(story.id),
                    "type": card.type,
                    "chapter": card.chapter,
                    "importance": card.importance,
                    "entities": ",".join(card.entities),
                    "tags": ",".join(card.tags),
                }
                for card in cards
            ],
            [card.id if card.id.startswith(f"{story.id}:") else f"{story.id}:memory_card:{card.id}" for card in cards],
        )
        self._audit_processed_chapter(story, chapter)

    def _audit_processed_chapter(self, story: Story, chapter: Chapter) -> None:
        """对已处理章节进行连续性审计，将报告写入故事对象的 continuity_reports 中。"""
        outline = None
        try:
            outline = story.get_outline(chapter.index)
        except KeyError:
            outline = None
        query = outline.summary if outline else chapter.summary or chapter.title
        longform_context = self.longform_manager.get_enhanced_context(chapter.index, story, query=query)
        report = self.continuity_auditor.audit_chapter(story, chapter.index, chapter.content, longform_context)
        story.continuity_reports[chapter.index] = report

    def _index_extracted_memory(self, story: Story, extraction) -> None:
        """将从章节中提取的角色、世界观和关系数据写入向量存储和图存储。"""
        characters = getattr(extraction, "characters", [])
        if characters:
            self.vector_store.add(
                "characters",
                [
                    " ".join(
                        part
                        for part in [
                            character.name,
                            str(character.age),
                            character.appearance,
                            character.personality,
                            character.motivation,
                            character.weakness,
                            character.arc,
                        ]
                        if part
                    )
                    for character in characters
                ],
                [{"story_id": str(story.id), "type": "character", "character_id": character.id} for character in characters],
                [f"{story.id}:character:{character.id}" for character in characters],
            )
            for character in characters:
                attrs = character.model_dump()
                attrs["story_id"] = str(story.id)
                self.graph_store.add_node(f"{story.id}:character:{character.id}", attrs)

        world_settings = getattr(extraction, "world_settings", [])
        if world_settings:
            self.vector_store.add(
                "world",
                [setting.content for setting in world_settings],
                [{"story_id": str(story.id), "type": "world", "category": setting.category, **setting.metadata} for setting in world_settings],
                [f"{story.id}:world:{setting.id}" for setting in world_settings],
            )

        for relation in getattr(extraction, "relationships", []):
            self.graph_store.add_edge(
                f"{story.id}:character:{relation.source}",
                f"{story.id}:character:{relation.target}",
                relation.relation,
            )

    def _require_story(self) -> Story:
        """返回当前活动故事，若未初始化则抛出 WorkflowError。"""
        if self.story is None:
            raise WorkflowError("No active story. Start or load a story first.")
        return self.story
